import sys

def extract_text_from_docx(file_path):
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed. Please try installing it.")
        sys.exit(1)
        
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())
            
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text.strip())
                    
    with open("scripts/resume_output.txt", "w", encoding="utf-8") as f:
        f.write('\n'.join(text))
    print("Successfully extracted docx to scripts/resume_output.txt")

if __name__ == "__main__":
    extract_text_from_docx(sys.argv[1])
